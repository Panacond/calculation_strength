def world(name, text_calc, loads):
    from docx import Document
    from docx.shared import Mm
    document = Document()
    if len(loads) > 2:

        document.add_heading('Таблица сбора нагрузок', level=1)

        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Вид нагрузки'
        hdr_cells[0].width=Mm(200)
        hdr_cells[1].text = 'Ед. изм.'
        hdr_cells[1].width=Mm(20)
        hdr_cells[1].autofit= False
        hdr_cells[2].text = 'Нормативная нагрузка'
        hdr_cells[2].width=Mm(20)
        hdr_cells[2].autofit= False
        hdr_cells[3].text = 'Коэффициент надежности по нагрузке'
        hdr_cells[3].width=Mm(20)
        hdr_cells[4].text = 'Расчетная нагрузка'
        hdr_cells[4].width=Mm(20)
        def select(a,i):
            if len(a)>1 or i == 0:
                return str(a[i])
            else:
                return ''
        for qty in loads[1:]:
            row_cells = table.add_row().cells
            row_cells[0].text = select(qty,0)
            row_cells[1].text = select(qty,1)
            row_cells[2].text = select(qty,2)
            row_cells[3].text = select(qty,3)
            row_cells[4].text = select(qty,4)
    
    import os
    if os.access(name +'1.png',os.F_OK):
        document.add_heading('Усилия', level=1)
        p = document.add_paragraph('Напряжения в элементе:')

        document.add_picture(name +'1.png' , width=Mm(160))
        document.add_picture(name +'2 Q.png' , width=Mm(160))
        document.add_picture(name +'3 M.png' , width=Mm(160))

        os.remove(name +'1.png')
        os.remove(name +'2 Q.png')
        os.remove(name +'3 M.png')
        p.add_run('').math = True

    document.add_heading('Расчетная часть', level=1)
    if '\n\n' in text_calc:
        text_calc = text_calc.replace('\n\n','\n')
    if '\t' in text_calc:
        text_calc = text_calc.replace('\t','    ')
    text_lisp = text_calc.split('\n')
    for i in text_lisp:
        document.add_paragraph(i)
    try:
        document.save(name +'.docx')
    except:
        print('Закрой файл ворлда!!!')

def main():
    name ='2Балка'
    text_calc= '\nM_r=q_r  \\times L^2/8=4000.0  \\times 6.0^2/8=18000.0'
    loads = [ ['1', '2', '3', '4', '5'], ['Плита над 2 этажом'], ['Постоянные '], ['ЖБ плита G=2500 кг/м.куб t=200 мм', 'кг/м.кв', 500, 1.1, 550], ['Итого', 'кг/м.кв', 500, '', 550], ['Временные'], ['Полезная G=200 кг/м.кв', 'кг/м.кв', 200, 1.2, 240], ['Итого', 'кг/м.кв', 200, '', 240], ['Всего', 'кг/м.кв', 700, '', 790], ['Всего временные', 'кг/м.кв', 200, '', 240], ['Всего постоянные ', 'кг/м.кв', 500, '', 550], ['Сумма всех нагрузок', 'кг/м.кв', 700, '', 790]]
    world(name, text_calc, loads )
    
if __name__ == '__main__':
    main()